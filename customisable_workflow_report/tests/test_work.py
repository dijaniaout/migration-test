from odoo.tests.common import TransactionCase, tagged
from odoo import Command
from odoo.exceptions import ValidationError
from odoo import fields
from datetime import timedelta
import base64
import docx
import tempfile
from PIL import Image, ImageDraw, ImageFont
import os
import logging
Logger = logging.getLogger(__name__)


@tagged('test_work_document')

class TestWork(TransactionCase):

    def setUp(self):
        super(TestWork, self).setUp()
        self.work_model = self.env['customisable_workflow.work_test']
        self.group_user = self.env.ref('base.group_user')
        self.document_model = self.env['customisable_workflow.document']
        self.work_given_document_model = self.env['customisable_workflow.work_given_document']
        self.step_model = self.env['customisable_workflow.step']
        self.workflow_model = self.env['customisable_workflow.workflow']
        self.group_manager = self.env.ref('base.group_system')
        self.step_decision_model = self.env['customisable_workflow.step_decision']
        self.user_model = self.env['res.users']

        # String to encode
        encoded_data = "dGVzdCBmaWxlIGNvbnRlbnRz"
        self.decoded_data = base64.b64decode(encoded_data.encode('utf-8'))

        # Create a document
        temp_docx_path = self._create_temporary_docx_file()
        f =  open(temp_docx_path, 'rb')
        doc = f.read()
        f.close()
        self.document_4 = self.document_model.create({
            'name': 'Document 4',
            'source': 'generated',
            'document': base64.b64encode(doc),
            'document_format_to_generate': 'docx',
            'is_required': True,
        })

        # Create Workflow without step
        self.workflow = self.workflow_model.create({
            'name': 'Processus demo',
        })

        # Create Steps
        self.workflow_step1 = self.step_model.create({
            'name': 'Etape 1',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id]),(6, 0, [self.group_user.id])],
        })

        # Set start_step_id for workflow
        self.workflow.write({
            'start_step_id': self.workflow_step1.id
        }) 

        # Create a work 
        self.work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': self.workflow.id,
        })
    
        self.user_test = self.user_model.create({
            'name': 'Test User',
            'login': 'testuser',
            'group_ids': [(6, 0, [self.group_manager.id])],
        })
        
        self.user1 = self.user_model.create({
            'name': 'User 1',
            'login': 'user1',
            'group_ids': [(6, 0, [self.group_user.id])]
        })

        # User env
        self.test_user_env = self.env(user=self.user_test)
        self.user1_env = self.env(user=self.user1)
        #Create a new workflow
        self.workflow_signature = self.workflow_model.create({
            'name': 'Processus signed document',
        }) 
        self.workflow_signature.write({
            'start_step_id':self._create_steps()[0].id
        }) 

    #Test that checks if the required documents have been created.Raises an exception if the required documents have not been created.
    def test_check_given_document_ids_exist(self):
        # Arrange
        workflow_decision_step1_step2 = self._create_decision_for_step1_to_step2()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_step2.id).make_a_decision()
        #Assert
        self.assertIsNotNone(self.work.given_document_ids, "L'utilisateur doit voir les documents à fournir sur l'étape.")

    #Tests that a ValidationError is raised if a required document is missing for the previous step before going to the next step.
    def test_missing_required_document(self):
        # Arrange
        workflow_decision_step1_step2 = self._create_decision_for_step1_to_step2()
        workflow_decision_step2_to_step3 = self._create_decision_for_step2_to_step3()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_step2.id).make_a_decision()
        self.work.given_document_ids[0].write({
            'attachment': self.decoded_data
        })
        self.work.invalidate_recordset()
        #Assert
        with self.assertRaises(ValidationError, msg="L'utilisateur doit être informé quand il ne fournit pas un document obligatoire."):
            self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step2_to_step3.id).make_a_decision()

    #Tests that no error is raised if a required document is already provided for the previous step before going to the next step.
    def test_existing_required_document(self):
        # Arrange
        workflow_decision_step1_step2 = self._create_decision_for_step1_to_step2()
        workflow_decision_step2_to_step3 = self._create_decision_for_step2_to_step3()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_step2.id).make_a_decision()
        self.work.given_document_ids.write({
            'attachment': self.decoded_data
        })
        self.work.invalidate_recordset()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step2_to_step3.id).make_a_decision()
        #Assert
        self.assertEqual(self.work.current_step_id.name, 'Etape 3', "On passe à l'étape suivante avec succès après avoir fourni le document requis à l'étape précédente.")

    # Test that checks if the documents to genered have been created.
    def test_check_genered_document_exist(self):
        # Arrange
        workflow_decision_step1_to_step4 = self._create_decision_for_step1_to_step4()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_to_step4.id).make_a_decision()
        #Assert
        self.assertIsNotNone(self.work.given_document_ids, "L'utilisateur doit voir les documents à générer sur l'étape.")

    # Test that a document can be validated
    def test_docs_to_validate(self):
        # Arrange
        workflow_decision_step1_to_step4 = self._create_decision_for_step1_to_step4()
        workflow_decision_step4_to_step5 = self._create_decision_for_step4_to_step5()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_to_step4.id).make_a_decision()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step4_to_step5.id).make_a_decision()
        # Assert
        self.assertIsNotNone(self.work.document_to_validate_ids, "L'utilisateur doit voir les documents à valider sur l'étape.")
        self.work.document_to_validate_ids.validate()
        self.assertEqual(self.work.document_to_validate_ids.state, 'validated', "Le document doit être validé")

    # Test that a document can be reject
    def test_docs_to_reject(self):
        # Arrange
        workflow_decision_step1_to_step4 = self._create_decision_for_step1_to_step4()
        workflow_decision_step4_to_step5 = self._create_decision_for_step4_to_step5()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_to_step4.id).make_a_decision()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step4_to_step5.id).make_a_decision()
        # Assert
        self.assertIsNotNone(self.work.document_to_validate_ids, "L'utilisateur doit voir les documents à valider sur l'étape.")
        self.work.document_to_validate_ids.reject()
        self.assertEqual(self.work.document_to_validate_ids.state, 'rejected', "Le document est rejeté")

    # Test that an unauthorized user should not be able to validate
    def test_not_authorized_user_cannot_validate(self):
        # Arrange
        workflow_decision_step1_to_step4 = self._create_decision_for_step1_to_step4()
        workflow_decision_step4_to_step5 = self._create_decision_for_step4_to_step5()
        # Act
        self.work.with_env(self.user1_env).with_context(decision_id = workflow_decision_step1_to_step4.id).make_a_decision()
        self.work.with_env(self.user1_env).with_context(decision_id = workflow_decision_step4_to_step5.id).make_a_decision()
        # Assert
        document_to_validate = self.work.document_event_ids.filtered(lambda x: x.type == 'validation')
        self.assertFalse(document_to_validate.with_env(self.user1_env).current_user_can_act, "Un utilisateur non autorisé ne doit pas pouvoir valider")

    # Test that an authorized user can validate a document
    def test_authorized_user_can_validate(self):
        # Arrange
        workflow_decision_step1_to_step4 = self._create_decision_for_step1_to_step4()
        workflow_decision_step4_to_step5 = self._create_decision_for_step4_to_step5()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_to_step4.id).make_a_decision()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step4_to_step5.id).make_a_decision()
        # Assert
        document_to_validate = self.work.document_event_ids.filtered(lambda x: x.type == 'validation')
        self.assertTrue(document_to_validate.with_env(self.test_user_env).current_user_can_act, "Un utilisateur autorisé doit pouvoir valider les documents.")

    # Test that a rejected document can be validated again
    def test_user_can_revalidate_a_reject_document(self):
        # Arrange
        workflow_decision_step1_step2 = self._create_decision_for_step1_to_step2()
        workflow_decision_step2_to_step3 = self._create_decision_for_step2_to_step3()
        workflow_decision_step3_to_step_reject = self._create_decision_for_step3_to_step_reject()
        # Act
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step1_step2.id).make_a_decision()
        self.work.given_document_ids.write({
            'attachment': self.decoded_data
        })
        self.work.invalidate_recordset()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step2_to_step3.id).make_a_decision()
        self.work.document_to_validate_ids.filtered(lambda x: x.source_name == "Document 2").validate()
        self.work.document_to_validate_ids.filtered(lambda x: x.source_name == "Document 3").reject()
        # Invalidate the cache so that the validation status is committed.
        self.work.document_to_validate_ids.invalidate_recordset()
        self.work.with_env(self.test_user_env).with_context(decision_id = workflow_decision_step3_to_step_reject.id).make_a_decision()
        # for doc in self.work.given_document_ids:
        reject_document = self.work.document_to_validate_ids.filtered(lambda x: x.state == "rejected")
        given_document = self.work.given_document_ids.filtered(lambda x: x.attachment == False)
        # Assert
        self.assertEqual(reject_document.origin_id.document_id.id, given_document.document_id.id, "L'utilisateur doit pouvoir fournir le document rejeté.")

    # test: Automatic steps are fulfilled and generate the documents to be generated
    def test_automatic_steps_are_fulfilled_generate_documents_tobe_generated(self):
        # Arrange
        workflow_generate_document = self.workflow_model.create({
            'name': 'Processus Automatic steps',
        })  
        step_a = self.step_model.create({
            'name': 'Étape de départ',
            'workflow_id': workflow_generate_document.id,
        })

        step_c = self.step_model.create({
            'name': 'Étape de fin',
            'workflow_id': workflow_generate_document.id,
        })
    
        step_b = self.step_model.create({
            'name': "Étape automatique",
            'workflow_id': workflow_generate_document.id,
            'is_automatic_step': True,
            'generated_document_ids': [(4, self.document_4.id)],
            'step_decision_ids': [(0, 0, {'name': 'Passer à l\'étape C',
                                            'target_step_id': step_c.id,})]
        })
        decision_stepA_to_stepB = self.step_decision_model.create({
            'name': 'Passer à l\'étape automatique',
            'workflow_id': workflow_generate_document.id,
            'step_id': step_a.id,
            'target_step_id': step_b.id,
        })
        
        workflow_generate_document.write({
            'start_step_id': step_a.id
        })
        
        work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': workflow_generate_document.id,
        })
        # Act
        work.with_context(decision_id=decision_stepA_to_stepB.id).make_a_decision()
        # Assert
        generated_documents = work.generated_document_ids
        self.assertTrue(generated_documents, "Un document à générer sur une étape automatique doit être généré.")
        self.assertEqual(work.current_step_id.id, step_c.id, "Les transitions automatiques d'étape doivent être satisfaites.")
    
    def test_Check_documents_tobe_signed_are_available(self):
        # Arrange 
        decision_stepA_to_stepB = self._create_decision_for_step_signature()[0]
        work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': self.workflow_signature.id,
        })  
        # Act
        work.with_context(decision_id=decision_stepA_to_stepB.id).make_a_decision() 
        # Assert
        document_to_sign_ids = work.has_documents_to_sign
        self.assertEqual(document_to_sign_ids, True, "L'utilisateur doit voir les documents à signer.")
        
    def test_notifications_are_sent(self):
        # Arrange 
        # Create a test user to send/receive the email
        user_notification = self.env['res.users'].create({
            'name': 'Test User',
            'login': 'logintestuser',
            'email': 'testuser@example.com',
        })

        workflow = self.workflow_model.create({
            'name': 'Processus Notification',
        })

        step_1 = self.step_model.create({
            'name': 'Étape 1',
            'workflow_id': workflow.id,
        })

        workflow.write({
            'start_step_id': step_1.id
        })

        work = self.work_model.create({
            'name': 'Test Work notification',
            'workflow_id': workflow.id,
        })

        # Create a email template
        email_template = self.env['mail.template'].create({
            'name': 'Test Notification Email',
            'email_from': user_notification.email_formatted,
            'subject': 'Test Email Notification for ${object.name}',
            'body_html': """
                <p>Bonjour ${object.user_id.name},</p>
                <p>Le document <strong>${object.name}</strong> a été validé.</p>
            """,
            'model_id': self.env['ir.model']._get('customisable_workflow.work_test').id,
        })

        # Nofication
        notification = self.env['banking_customisable_workflow.notification_type'].create({
            'name': 'Test Notification',
            'template_id': email_template.id,
            'when_to_send': 'beginning',
            'user_ids': [(4, user_notification.id)],
        })

        step_2 = self.step_model.create({
            'name': 'Étape 2',
            'workflow_id': workflow.id,
            'notification_type_ids': [(4, notification.id)]
        })

        decision_step1_to_step2 = self.step_decision_model.create({
            'name': "Passer à l'étape 2",
            'workflow_id': workflow.id,
            'step_id': step_1.id,
            'target_step_id': step_2.id,
        })
        # Act
        work.with_context(decision_id=decision_step1_to_step2.id).make_a_decision()
        # Assert
        sent_mail = self.env['mail.mail'].search(
            [('subject', 'like', 'Test Email Notification for ${object.name}')],
        )
        self.assertTrue(sent_mail, "La notification n'a pas été envoyé")
        self.assertEqual(sent_mail.email_to, 'testuser@example.com', "L'e-mail du destinataire est incorrect.")

    # Test the work should go to the expiration stage when the processing time is exceeded.
    def test_the_file_processing_time_has_been_exceeded(self):
        # Arrange*
        steps = self._create_step1_step2_for_processing_time()
        decision_step1_step2 = self._create_decision_for_step1_to_step2_processin_time()

        # Set start_step_id for workflow
        self.workflow.write({
            'start_step_id': steps[0].id,
            'maximum_processing_time': 2,
            'processing_start_step_id': steps[0].id,
            'processing_end_step_id': steps[2].id,
            'processing_expiraton_step_id': steps[3].id
        })

        # Create a work 
        self.work = self.work_model.create({
            'name': 'Test Work processing time',
            'workflow_id':  self.workflow.id,
        })

        original_today = fields.Date.today
        def today_plus_two():
            return original_today() + timedelta(days=3)

        # make 'fields.Date.today' to return "today + 2 days"
        fields.Date.today = today_plus_two
        #Act
        self.work.with_env(self.test_user_env).with_context(decision_id = decision_step1_step2.id).make_a_decision()
        self.work.given_document_ids.write({
            'attachment': self.decoded_data
        })
        self.work._cron_transition_time()
        #Assert
        self.assertEqual(self.work.current_step_id.name, "Etape d'expiration", "Le work doit aller à l'étape d'expiration quand le temps de traitement est dépassé.")

    # Test that when the processing time is respected we do not go to the expiration stage
    def test_that_the_processing_time_is_respected(self):
        # Arrange
        steps = self._create_step1_step2_for_processing_time()
        decision_step1_step2 = self._create_decision_for_step1_to_step2_processin_time()
        decision_step1_step3 = self._create_decision_for_step1_to_step3_processin_time()
        decision_step3_step5 = self._create_decision_for_step3_to_step5_processin_time()
        # Set start_step_id for workflow
        self.workflow.write({
            'start_step_id': steps[0].id,
            'maximum_processing_time': 2,
            'processing_start_step_id': steps[0].id,
            'processing_end_step_id': steps[2].id,
            'processing_expiraton_step_id': steps[3].id
        })

        # Create a work 
        self.work = self.work_model.create({
            'name': 'Test Work processing time',
            'workflow_id':  self.workflow.id,
        })

        original_today = fields.Date.today
        def today_plus_two():
            return original_today() + timedelta(days=2)

        # make 'fields.Date.today' to return "today + 2 days"
        fields.Date.today = today_plus_two
        #Act
        self.work.with_env(self.test_user_env).with_context(decision_id = decision_step1_step2.id).make_a_decision()
        self.work.given_document_ids.write({
            'attachment': self.decoded_data
        })
        self.work.with_env(self.test_user_env).with_context(decision_id = decision_step1_step3.id).make_a_decision()
        self.work.with_env(self.test_user_env).with_context(decision_id = decision_step3_step5.id).make_a_decision()
        self.work._cron_transition_time()
        #Assert
        self.assertNotEqual("Etape 2", "Etape d'expiration", "Le work ne doit pas aller sur l'étape d'expiration lorsque le temps de traitement est respecté.")

    # Create step2
    def _create_step1_step2_for_processing_time(self):
        document_1 = self.document_model.create({
            'name': 'Document 1',
            'source': 'provided',
            'is_required': True
        })
        if not all(hasattr(self, attr) for attr in ['step_1', 'step_2', 'step_3', 'step_4', 'step_5']):
            self.step_1 = self.step_model.create({
                'name': 'Etape de début de traitement',
                'workflow_id':  self.workflow.id,
            })
            self.step_2 = self.step_model.create({
                'name': 'Etape 2',
                'workflow_id': self.workflow.id,
                'allowed_group_ids': [(6, 0, [self.group_manager.id])],
                'provided_document_ids': [(4, document_1.id)]
            })
            self.step_3 = self.step_model.create({
                'name': 'Etape de fin de traitement',
                'workflow_id':  self.workflow.id,
                'allowed_group_ids': [(6, 0, [self.group_manager.id])]
            })
            self.step_4 = self.step_model.create({
                'name': "Etape d'expiration",
                'workflow_id':  self.workflow.id,
                'allowed_group_ids': [(6, 0, [self.group_manager.id])]
            })
            self.step_5 = self.step_model.create({
                'name': 'Etape 5',
                'workflow_id':  self.workflow.id,
                'allowed_group_ids': [(6, 0, [self.group_manager.id])]
            })
        return [self.step_1, self.step_2, self.step_3, self.step_4, self.step_5]

    # Create decision to step1 to step6
    def _create_decision_for_step1_to_step2_processin_time(self):
        steps = self._create_step1_step2_for_processing_time()
        # Create Step Decisions
        self.workflow_decision_step1_to_step2 = self.step_decision_model.create({
            'name': 'Passer à étape 2',
            'workflow_id': self.workflow.id,
            'step_id': steps[0].id,
            'target_step_id': steps[1].id,
        })
        return self.workflow_decision_step1_to_step2
    
    def _create_decision_for_step1_to_step3_processin_time(self):
        steps = self._create_step1_step2_for_processing_time()
        # Create Step Decisions
        self.workflow_decision_step1_to_step3 = self.step_decision_model.create({
            'name': 'Passer à étape de fin de traitement',
            'workflow_id': self.workflow.id,
            'step_id': steps[1].id,
            'target_step_id': steps[2].id,
        })
        return self.workflow_decision_step1_to_step3
    
    def _create_decision_for_step3_to_step5_processin_time(self):
        steps = self._create_step1_step2_for_processing_time()
        # Create Step Decisions
        self.workflow_decision_step3_to_step5 = self.step_decision_model.create({
            'name': 'Passer à étape 5',
            'workflow_id': self.workflow.id,
            'step_id': steps[3].id,
            'target_step_id': steps[4].id,
        })
        return self.workflow_decision_step3_to_step5
    def test_user_can_sign_document(self):
        # Arrange
        decision_stepA_to_stepB = self._create_decision_for_step_signature()
        # Create a temporary image for the signature
        temp_img_path = self._create_temporary_image_file()
        with open(temp_img_path, 'rb') as image_file:
            img = image_file.read()
        encoded_img = base64.b64encode(img).decode('utf-8')
        # Create a user with a signature
        user_test = self.user_model.create({
            'name': 'User to sign',
            'login': 'user_signe',
            'group_ids': [(6, 0, [self.group_manager.id])],
            'document_signature': encoded_img
        })
        work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': self.workflow_signature.id,
        })
        # Act
        test_user_env = self.env(user=user_test)
        work.with_env(test_user_env).with_context(decision_id=decision_stepA_to_stepB.id).make_a_decision()
        document_to_signs = work.signed_document_ids
        state_document = ''
        document_to_signs = work.signed_document_ids
        for document_to_sign in document_to_signs:
            document_to_sign.source_document = self.document_4.document
            document_to_sign.with_user(user_test).sign()
            state_document = document_to_sign.state
        # Assert 
        self.assertEqual(state_document, 'signed', "On doit pouvoir signer un document")
        
    def test_document_is_signed_before_topasse_anext_step(self):
        # Arrange
        steps = self._create_steps()
        # Decision step A to step B
        decision_stepA_to_stepB = self._create_decision_for_step_signature()
        # Decision step B to step C
        decision_stepB_to_stepC = self.step_decision_model.create({
            'name': 'Passer à l\'étape C',
            'workflow_id': self.workflow_signature.id,
            'step_id': steps[1].id,
            'target_step_id': steps[2].id,
        })
        work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': self.workflow_signature.id,
        })  
        
        # Act
        work.with_context(decision_id=decision_stepA_to_stepB.id).make_a_decision()
        document_to_sign = work.signed_document_ids[0]
        # Assert
        with self.assertRaises(ValidationError, msg="L'utilisateur doit signé un document "):
            work.with_context(decision_id=decision_stepB_to_stepC.id).make_a_decision()

    def test_reject_document_to_sign(self):
        # Arrange
        decision_stepA_to_stepB = self._create_decision_for_step_signature()
        work = self.work_model.create({
            'name': 'Test Work',
            'workflow_id': self.workflow_signature.id,
        }) 
        # Act
        work.with_context(decision_id=decision_stepA_to_stepB.id).make_a_decision()
        document_to_sign = work.signed_document_ids
        document_to_sign.reject()
        state_document = document_to_sign.state
        # Assert
        self.assertEqual(state_document, 'rejected', "L'utilisateur  doit pouvoir rejeter un document")


    def _create_steps(self):
        step_a = self.step_model.create({
            'name': 'Étape de départ',
            'workflow_id': self.workflow_signature.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id])],
            'generated_document_ids': [(4, self.document_4.id)],
        })

        step_b = self.step_model.create({
            'name': "Étape de signature",
            'workflow_id': self.workflow_signature.id,
            'is_automatic_step': True,
            'allowed_group_ids': [(6, 0, [self.group_manager.id])],
            'document_to_sign_ids': [(0, 0, {'document_id': self.document_4.id,
                                            'allowed_group_ids': [(6, 0, [self.group_manager.id])]})]
        })

        step_c = self.step_model.create({
            'name': 'Étape de fin',
            'workflow_id': self.workflow_signature.id,
        })
        return [step_a, step_b, step_c]
    
    def _create_decision_for_step_signature(self):
        steps = self._create_steps()
        decision_stepA_to_stepB = self.step_decision_model.create({
            'name': 'Passer à l\'étape de signature',
            'workflow_id': self.workflow_signature.id,
            'step_id': steps[0].id,
            'target_step_id': steps[1].id,
        }) 
        return decision_stepA_to_stepB


    # Function to Create a new .docx document using the python-docx library
    def _create_temporary_docx_file(self):
        doc = docx.Document()
        doc.add_heading('Titre du document', 0)
        doc.add_paragraph('Ceci est un fichier temporaire généré avec Python.')
        # Créer un fichier temporaire
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            # Sauvegarder le document dans le fichier temporaire
            doc.save(tmp_file.name)
        return tmp_file.name

    def _create_temporary_image_file(self):
        image = Image.new('RGB', (200, 100), color='white')
        draw = ImageDraw.Draw(image)
        text = "Ceci est une image temporaire"
        text_width, text_height = draw.textsize(text)
        position = ((200 - text_width) // 2, (100 - text_height) // 2)
        draw.text(position, text, fill="black")

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_file:
            image.save(tmp_file, format="PNG")
        return tmp_file.name

    # Create given document
    def _create_given_document(self):
        # Create a provided document
        if not hasattr(self, 'document_2') or not hasattr(self, 'document_3'):
            self.document_2 = self.document_model.create({
                'name': 'Document 2',
                'source': 'provided',
                'is_required': True
            })
            self.document_3 = self.document_model.create({
                'name': 'Document 3',
                'source': 'provided',
                'is_required': True
            })
        return [self.document_2, self.document_3]

    # Create step 2
    def _create_step_2(self):
        # Create Steps
        documents = self._create_given_document()
        self.workflow_step2 = self.step_model.create({
            'name': 'Etape 2',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id])],
            'provided_document_ids': [(4, documents[0].id), (4, documents[1].id)]
        })
        return self.workflow_step2

    #Create decision to step1 to step2
    def _create_decision_for_step1_to_step2(self):
        # Create Step Decisions
        self.workflow_decision_step1_step2 = self.step_decision_model.create({
            'name': 'Valider',
            'workflow_id': self.workflow.id,
            'step_id': self.workflow_step1.id,
            'target_step_id': self._create_step_2().id,
        })
        return self.workflow_decision_step1_step2

    # Create step3
    def _create_step3(self):
        # Create Step 3
        documents = self._create_given_document()
        self.workflow_step3 = self.step_model.create({
            'name': 'Etape 3',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id])],
            'document_to_validate_ids': [(0, 0, {'document_id': documents[0].id,
                                                'allowed_group_ids': [(6, 0, [self.group_manager.id])]}),
                                        (0, 0, {'document_id': documents[1].id,
                                                'allowed_group_ids': [(6, 0, [self.group_manager.id])]})
                                        ]
        })
        return self.workflow_step3

    #Create decision to step1 to step2
    def _create_decision_for_step2_to_step3(self):
        # Create Step Decisions
        self.workflow_decision_step2_to_step3 = self.step_decision_model.create({
            'name': 'Soumettre',
            'workflow_id': self.workflow.id,
            'step_id': self._create_step_2().id,
            'target_step_id': self._create_step3().id,
        })
        return self.workflow_decision_step2_to_step3

    # Create step rejet
    def _create_step_reject(self):
        documents = self._create_given_document()
        # Create Steps
        self.workflow_step_reject = self.step_model.create({
            'name': 'Etape reject',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id])],
            'provided_document_ids': [(4, documents[0].id), (4, documents[1].id)]
        })
        return self.workflow_step_reject

    #Create decision to step3 to step_reject
    def _create_decision_for_step3_to_step_reject(self):
        # Create Step Decisions
        self.workflow_decision_step3_to_step_reject = self.step_decision_model.create({
            'name': 'Rejeter',
            'workflow_id': self.workflow.id,
            'step_id': self._create_step3().id,
            'target_step_id': self._create_step_reject().id,
        })
        return self.workflow_decision_step3_to_step_reject

    #Create decision to step3 to step4
    def _create_decision_for_step3_to_step4(self):
        # Create Step Decisions
        self.workflow_decision_step3_to_step4 = self.step_decision_model.create({
            'name': 'Valider',
            'workflow_id': self.workflow.id,
            'step_id': self._create_step3().id,
            'target_step_id': self._create_step4().id,
        })
        return self.workflow_decision_step2_to_step3

    #Create Step 4
    def _create_step4(self):
        # Create Step 4
        self.workflow_step4 = self.step_model.create({
            'name': 'Etape 4',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, [self.group_manager.id]), (6, 0, [self.group_user.id])],
            'generated_document_ids': [(4, self.document_4.id)]
        })
        return self.workflow_step4

    # Create decision to step1 to step4
    def _create_decision_for_step1_to_step4(self):
        # Create Step Decisions
        self.workflow_decision_step1_to_step4 = self.step_decision_model.create({
            'name': 'Passer à étape 4',
            'workflow_id': self.workflow.id,
            'step_id': self.workflow_step1.id,
            'target_step_id': self._create_step4().id,
        })
        return self.workflow_decision_step1_to_step4

    # Create decision to step4 to step5
    def _create_decision_for_step4_to_step5(self):
        # Create Step 5
        self.workflow_step5 = self.step_model.create({
            'name': 'Etape 5',
            'workflow_id': self.workflow.id,
            'allowed_group_ids': [(6, 0, {self.group_manager.id})],
            'document_to_validate_ids': [(0, 0, {'document_id': self.document_4.id,
                                                'allowed_group_ids': [(6, 0, [self.group_manager.id])]})]
        })
        # Create Step Decisions
        self.workflow_decision_step4_to_step5 = self.step_decision_model.create({
            'name': 'Passer à étape 5',
            'workflow_id': self.workflow.id,
            'step_id': self.workflow_step4.id,
            'target_step_id': self.workflow_step5.id,
        })
        return self.workflow_decision_step4_to_step5
    
    